import os
import markdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn
from bs4 import BeautifulSoup, NavigableString
from typing import Optional, List, Dict
import re

def convert_md_to_docx_pure_python(input_file: str, output_file: Optional[str] = None) -> str:
    """
    使用纯 Python 库将 Markdown 转换为 Word 文档
    """
    if not os.path.exists(input_file):
        print(f"[错误] 输入文件不存在: {input_file}")
        return None
    
    if output_file is None:
        output_file = os.path.splitext(input_file)[0] + '.docx'
    
    try:
        # 读取 Markdown 文件
        with open(input_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 预处理：处理参考文献格式
        md_content = preprocess_markdown(md_content)
        
        # 转换为 HTML
        html = markdown.markdown(md_content, extensions=[
            'tables', 
            'fenced_code', 
            'codehilite',
            'nl2br',
            'toc'  # 添加目录支持
        ])
        
        # 解析 HTML
        soup = BeautifulSoup(html, 'html.parser')
        
        # 创建 Word 文档
        doc = Document()
        
        # 设置文档样式（支持中文字体）
        setup_chinese_document_styles(doc)
        
        # 提取标题信息用于后续目录生成
        headings = extract_headings(soup)
        
        # 按顺序处理所有元素，在摘要后插入目录
        process_elements_with_toc_insertion(doc, soup, headings)
        
        # 保存文档
        doc.save(output_file)
        print(f"✅ Word 文档已生成: {output_file}")
        return output_file
        
    except Exception as e:
        print(f"[错误] 转换失败: {e}")
        import traceback
        traceback.print_exc()
        return None

def process_elements_with_toc_insertion(doc, soup, headings):
    """
    按顺序处理HTML元素，在摘要后插入目录
    """
    body = soup.find('body') if soup.find('body') else soup
    toc_inserted = False
    skip_markdown_toc = False
    
    for element in body.children:
        if hasattr(element, 'name') and element.name:
            # 跳过 Markdown 中的目录部分
            if element.name == 'h2' and '目录' in element.get_text():
                skip_markdown_toc = True
                continue
            
            if skip_markdown_toc and element.name in ['ul', 'ol', 'li']:
                continue
                
            if skip_markdown_toc and element.name and element.name.startswith('h'):
                skip_markdown_toc = False
            
            # 处理当前元素
            if not skip_markdown_toc:
                process_single_element(doc, element)
                
                # 检查是否刚处理完摘要，如果是则插入目录
                if not toc_inserted and element.name and element.name.startswith('h'):
                    text = clean_text_with_references(element.get_text())
                    if '摘要' in text or 'Abstract' in text:
                        # 在摘要后插入目录
                        if headings:
                            doc.add_page_break()  # 摘要后分页
                            add_improved_table_of_contents(doc, headings)
                            doc.add_page_break()  # 目录后分页
                            toc_inserted = True

def add_improved_table_of_contents(doc, headings: List[Dict]):
    """
    添加改进的目录到文档
    """
    # 添加目录标题
    toc_heading = doc.add_heading('目录', level=1)
    toc_heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    for run in toc_heading.runs:
        run.font.name = '黑体'
        run.font.size = Pt(16)
        run.font.bold = True
        set_run_chinese_font(run, '黑体')
    
    # 添加空行
    doc.add_paragraph()
    
    # 添加目录项
    for i, heading in enumerate(headings):
        level = heading['level']
        text = heading['text']
        
        # 跳过目录标题本身
        if '目录' in text:
            continue
            
        # 根据标题级别设置缩进和样式
        para = doc.add_paragraph()
        
        # 设置缩进
        if level == 1:
            para.paragraph_format.left_indent = Inches(0)
            font_size = Pt(14)
            is_bold = True
        elif level == 2:
            para.paragraph_format.left_indent = Inches(0.3)
            font_size = Pt(12)
            is_bold = False
        else:
            para.paragraph_format.left_indent = Inches(0.6)
            font_size = Pt(11)
            is_bold = False
        
        # 设置行间距
        para.paragraph_format.space_after = Pt(6)
        
        # 添加标题文本
        run = para.add_run(text)
        run.font.name = '宋体'
        run.font.size = font_size
        run.font.bold = is_bold
        set_run_chinese_font(run, '宋体')
        
        # 添加制表符和页码占位符
        tab_run = para.add_run('\t')
        
        # 添加点线
        dots_run = para.add_run('.' * 50)
        dots_run.font.name = '宋体'
        dots_run.font.size = Pt(10)
        set_run_chinese_font(dots_run, '宋体')
        
        # 添加页码占位符
        page_run = para.add_run(f'\t{i + 1}')
        page_run.font.name = '宋体'
        page_run.font.size = font_size
        page_run.font.bold = is_bold
        set_run_chinese_font(page_run, '宋体')
        
        # 设置制表位（右对齐页码）
        tab_stops = para.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Inches(6.0))
    
    # 添加空行
    doc.add_paragraph()

def preprocess_markdown(content: str) -> str:
    """
    预处理 Markdown 内容，特别处理参考文献
    """
    # 保护参考文献格式，避免被markdown解析器破坏
    content = re.sub(r'\[(\d+)\]', r'【参考文献\1】', content)
    return content

def extract_headings(soup) -> List[Dict]:
    """
    提取文档中的所有标题，用于生成目录
    """
    headings = []
    heading_tags = soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
    
    for heading in heading_tags:
        level = int(heading.name[1])
        text = clean_text_with_references(heading.get_text())
        if text.strip():
            headings.append({
                'level': level,
                'text': text.strip(),
                'id': f"heading_{len(headings)}"
            })
    
    return headings

def add_table_of_contents(doc, headings: List[Dict]):
    """
    添加目录到文档
    """
    # 添加目录标题
    toc_heading = doc.add_heading('目录', level=1)
    for run in toc_heading.runs:
        run.font.name = '黑体'
        set_run_chinese_font(run, '黑体')
    
    # 添加目录项
    for heading in headings:
        level = heading['level']
        text = heading['text']
        
        # 根据标题级别设置缩进
        indent_level = (level - 1) * 0.5  # 每级缩进0.5英寸
        
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(indent_level)
        
        # 添加标题文本
        run = para.add_run(text)
        run.font.name = '宋体'
        run.font.size = Pt(12 - level)  # 根据级别调整字体大小
        set_run_chinese_font(run, '宋体')
        
        # 添加页码占位符（实际Word中需要手动更新域）
        para.add_run('\t')
        page_run = para.add_run('...')
        page_run.font.name = '宋体'
        set_run_chinese_font(page_run, '宋体')
    
    # 添加空行
    doc.add_paragraph()

def setup_chinese_document_styles(doc):
    """
    设置支持简体中文的文档样式
    """
    styles = doc.styles
    
    # 设置正文样式
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = '宋体'
    normal_font.size = Pt(12)
    set_chinese_font(normal_style, '宋体')
    
    # 设置标题样式
    for i in range(1, 7):
        heading_style = styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.name = '黑体'
        heading_font.size = Pt(18 - i)
        set_chinese_font(heading_style, '黑体')

def set_chinese_font(style, font_name):
    """
    设置中文字体，确保中文字符正确显示
    """
    try:
        rPr = style._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:cs'), font_name)
        
    except Exception as e:
        print(f"设置中文字体时出错: {e}")

def process_elements_in_order_skip_toc(doc, soup):
    """
    按顺序处理HTML元素，但跳过目录部分
    """
    body = soup.find('body') if soup.find('body') else soup
    skip_toc = False
    
    for element in body.children:
        if hasattr(element, 'name') and element.name:
            # 检查是否是目录标题
            if element.name == 'h2' and '目录' in element.get_text():
                skip_toc = True
                continue
            
            # 检查是否是下一个主要章节（结束目录跳过）
            if skip_toc and element.name and element.name.startswith('h') and element.name != 'h2':
                skip_toc = False
            
            # 跳过目录中的列表项
            if skip_toc and element.name in ['ul', 'ol', 'li']:
                continue
                
            if not skip_toc:
                process_single_element(doc, element)

def process_elements_in_order(doc, soup):
    """
    按顺序处理HTML元素
    """
    body = soup.find('body') if soup.find('body') else soup
    
    for element in body.children:
        if hasattr(element, 'name') and element.name:
            process_single_element(doc, element)

def process_single_element(doc, element):
    """
    处理单个HTML元素
    """
    if element.name and element.name.startswith('h'):
        # 标题
        level = int(element.name[1])
        text = clean_text_with_references(element.get_text(), is_table_context=False)
        if text.strip():
            heading = doc.add_heading(text, level=level)
            for run in heading.runs:
                run.font.name = '黑体'
                set_run_chinese_font(run, '黑体')
            
    elif element.name == 'p':
        # 段落 - 检查是否包含表格相关内容
        text = element.get_text()
        is_table_related = '数据来源' in text or '基于' in text or '注：' in text
        text = clean_text_with_references(text, is_table_context=is_table_related)
        if text.strip():
            para = doc.add_paragraph(text)
            for run in para.runs:
                run.font.name = '宋体'
                set_run_chinese_font(run, '宋体')
            
            # 如果是表格描述段落，添加额外空行
            if '基于' in text and '数据来源' in text:
                doc.add_paragraph()  # 在表格描述后添加空行
            
    elif element.name == 'table':
        # 表格 - 使用改进的表格处理
        add_robust_table_to_doc(doc, element)
        
    elif element.name in ['ul', 'ol']:
        # 列表
        add_improved_list_to_doc(doc, element)
        
    elif element.name == 'pre':
        # 代码块
        add_code_block_to_doc(doc, element)
        
    elif element.name == 'blockquote':
        # 引用块
        text = clean_text_with_references(element.get_text())
        if text.strip():
            para = doc.add_paragraph(text)
            para.style = 'Quote'
            for run in para.runs:
                run.font.name = '楷体'
                set_run_chinese_font(run, '楷体')

def set_run_chinese_font(run, font_name):
    """
    为run设置中文字体
    """
    try:
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rFonts.set(qn('w:eastAsia'), font_name)
        rFonts.set(qn('w:cs'), font_name)
    except Exception as e:
        print(f"设置run中文字体时出错: {e}")

def clean_text_with_references(text: str, is_table_context: bool = False) -> str:
    """
    清理文本内容并正确处理参考文献，过滤数据来源信息
    """
    text = re.sub(r'\s+', ' ', text)
    # 修复：将【参考文献1】转换回正确的[1]格式
    text = re.sub(r'【参考文献(\d+)】', r'[\1]', text)
    
    # 如果是表格上下文，则保留数据来源信息
    if not is_table_context:
        # 过滤数据来源信息
        # 匹配各种数据来源格式并移除
        text = re.sub(r'\(数据来源[：:][^)]*\)', '', text)  # (数据来源：xxx)
        text = re.sub(r'数据来源[：:][^，。；\n]*[，。；]?', '', text)  # 数据来源：xxx
        text = re.sub(r'来源[：:][^，。；\n]*[，。；]?', '', text)  # 来源：xxx
        text = re.sub(r'资料来源[：:][^，。；\n]*[，。；]?', '', text)  # 资料来源：xxx
        text = re.sub(r'数据源[：:][^，。；\n]*[，。；]?', '', text)  # 数据源：xxx
        text = re.sub(r'\([^)]*来源[^)]*\)', '', text)  # (xxx来源xxx)
        text = re.sub(r'\([^)]*数据[^)]*\)', '', text)  # (xxx数据xxx)
    
    # 清理多余的空格和标点
    text = re.sub(r'\s+', ' ', text)
    text = re.sub(r'^[，。；：\s]+|[，。；：\s]+$', '', text)  # 移除开头结尾的标点和空格
    
    return text.strip()

def add_robust_table_to_doc(doc, table_element):
    """
    强化的表格处理，解决表格转换错误问题并美化样式
    """
    try:
        # 在表格前添加空行，确保与前面内容有适当间距
        doc.add_paragraph()
        
        # 获取所有行
        all_rows = table_element.find_all('tr')
        if not all_rows:
            print("[警告] 表格中没有找到行")
            return
        
        # 过滤掉包含数据来源信息的行
        rows = []
        source_info = None
        
        for row in all_rows:
            row_text = row.get_text().strip()
            # 检查是否是数据来源行
            if '数据来源' in row_text or '来源' in row_text:
                source_info = clean_text_with_references(row_text, is_table_context=True)
                continue
            rows.append(row)
        
        if not rows:
            print("[警告] 过滤后表格中没有有效行")
            return
        
        # 分析表格结构
        table_data = []
        max_cols = 0
        
        for row_idx, row in enumerate(rows):
            # 获取单元格（包括th和td）
            cells = row.find_all(['th', 'td'])
            row_data = []
            
            for cell in cells:
                cell_text = clean_text_with_references(cell.get_text(), is_table_context=True)
                # 处理合并单元格
                colspan = int(cell.get('colspan', 1))
                rowspan = int(cell.get('rowspan', 1))
                is_header = cell.name == 'th' or row_idx == 0
                
                # 添加单元格内容
                row_data.append({
                    'text': cell_text,
                    'colspan': colspan,
                    'rowspan': rowspan,
                    'is_header': is_header
                })
                
                # 如果有列合并，添加空单元格占位
                for _ in range(colspan - 1):
                    row_data.append({'text': '', 'colspan': 1, 'rowspan': 1, 'is_header': is_header})
            
            table_data.append(row_data)
            max_cols = max(max_cols, len(row_data))
        
        if max_cols == 0:
            print("[警告] 表格没有有效列")
            return
        
        # 标准化表格数据（确保所有行都有相同的列数）
        for row_data in table_data:
            while len(row_data) < max_cols:
                row_data.append({'text': '', 'colspan': 1, 'rowspan': 1, 'is_header': False})
        
        # 创建Word表格
        word_table = doc.add_table(rows=len(table_data), cols=max_cols)
        
        # 设置表格样式
        word_table.style = 'Light Grid Accent 1'  # 使用更美观的样式
        word_table.autofit = False
        
        # 设置表格宽度
        word_table.width = Inches(6.5)
        
        # 填充表格内容并设置样式
        for row_idx, row_data in enumerate(table_data):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < max_cols and row_idx < len(table_data):
                    try:
                        cell = word_table.cell(row_idx, col_idx)
                        cell.text = cell_data['text']
                        
                        # 设置单元格宽度（根据列数平均分配）
                        cell.width = Inches(6.5 / max_cols)
                        
                        # 设置单元格样式
                        for paragraph in cell.paragraphs:
                            # 设置段落对齐方式
                            if cell_data['is_header']:
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            elif col_idx == 0:  # 第一列左对齐
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                            else:  # 数据列居中对齐
                                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                            
                            # 设置字体样式
                            for run in paragraph.runs:
                                if cell_data['is_header']:
                                    # 表头样式：加粗、稍大字体
                                    run.font.name = '微软雅黑'
                                    run.font.size = Pt(11)
                                    run.font.bold = True
                                    set_run_chinese_font(run, '微软雅黑')
                                else:
                                    # 数据行样式：普通字体
                                    run.font.name = '宋体'
                                    run.font.size = Pt(10)
                                    run.font.bold = False
                                    set_run_chinese_font(run, '宋体')
                        
                        # 如果单元格为空，确保至少有一个段落
                        if not cell.text.strip() and not cell.paragraphs[0].runs:
                            run = cell.paragraphs[0].add_run('')
                            if cell_data['is_header']:
                                run.font.name = '微软雅黑'
                                run.font.bold = True
                                set_run_chinese_font(run, '微软雅黑')
                            else:
                                run.font.name = '宋体'
                                set_run_chinese_font(run, '宋体')
                        
                        # 设置单元格内边距
                        tc = cell._tc
                        tcPr = tc.get_or_add_tcPr()
                        tcMar = OxmlElement('w:tcMar')
                        
                        # 设置上下左右边距
                        for margin_name in ['top', 'left', 'bottom', 'right']:
                            margin = OxmlElement(f'w:{margin_name}')
                            margin.set(qn('w:w'), '100')  # 设置边距为100 twips (约1.4mm)
                            margin.set(qn('w:type'), 'dxa')
                            tcMar.append(margin)
                        
                        tcPr.append(tcMar)
                            
                    except Exception as cell_error:
                        print(f"[警告] 处理单元格 ({row_idx}, {col_idx}) 时出错: {cell_error}")
        
        # 设置表格边框
        try:
            tbl = word_table._tbl
            tblPr = tbl.tblPr
            
            # 设置表格边框
            tblBorders = OxmlElement('w:tblBorders')
            
            # 定义边框样式
            border_attrs = {
                'w:val': 'single',
                'w:sz': '4',  # 边框粗细
                'w:space': '0',
                'w:color': '000000'  # 黑色边框
            }
            
            # 添加各种边框
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                for attr, value in border_attrs.items():
                    border.set(attr, value)
                tblBorders.append(border)
            
            tblPr.append(tblBorders)
        except Exception as border_error:
            print(f"[警告] 设置表格边框时出错: {border_error}")
        
        # 添加表格后的空行 - 增加更多空行以便后续注释
        doc.add_paragraph()
        doc.add_paragraph()  # 额外的空行
        # 如果有数据来源信息，作为独立段落添加
        if source_info:
            source_para = doc.add_paragraph(source_info)
            for run in source_para.runs:
                run.font.name = '宋体'
                run.font.size = Pt(10)
                run.font.italic = True  # 斜体显示数据来源
                set_run_chinese_font(run, '宋体')
            # 数据来源后再添加空行
            doc.add_paragraph()
        
        print(f"✅ 成功添加美化表格: {len(table_data)}行 x {max_cols}列")
        
    except Exception as e:
        print(f"[错误] 表格处理失败: {e}")
        # 如果表格处理失败，尝试以文本形式添加
        try:
            table_text = table_element.get_text()
            if table_text.strip():
                para = doc.add_paragraph(f"[表格内容]: {table_text}")
                for run in para.runs:
                    run.font.name = '宋体'
                    set_run_chinese_font(run, '宋体')
        except:
            print("[错误] 表格文本备用方案也失败")

def add_improved_list_to_doc(doc, list_element):
    """
    改进的列表处理，支持中文字体
    """
    items = list_element.find_all('li', recursive=False)
    is_ordered = list_element.name == 'ol'
    
    for i, item in enumerate(items):
        # 处理列表项文本
        text_parts = []
        for content in item.children:
            if isinstance(content, NavigableString):
                text_parts.append(str(content).strip())
            elif hasattr(content, 'get_text'):
                text_parts.append(content.get_text().strip())
        
        text = ' '.join(filter(None, text_parts))
        text = clean_text_with_references(text)
        
        if text:
            if is_ordered:
                para = doc.add_paragraph(f"{i+1}. {text}")
            else:
                para = doc.add_paragraph(text, style='List Bullet')
            
            # 设置列表字体
            for run in para.runs:
                run.font.name = '宋体'
                set_run_chinese_font(run, '宋体')
        
        # 处理嵌套列表
        nested_lists = item.find_all(['ul', 'ol'], recursive=False)
        for nested_list in nested_lists:
            add_improved_list_to_doc(doc, nested_list)

def add_code_block_to_doc(doc, pre_element):
    """
    添加代码块，使用等宽字体
    """
    code_text = pre_element.get_text()
    if code_text.strip():
        para = doc.add_paragraph(code_text)
        # 设置等宽字体
        for run in para.runs:
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
            set_run_chinese_font(run, 'Consolas')

if __name__ == "__main__":
    # 测试转换
    test_files = [
        "/Users/mac/Documents/ai/bishai/financial_research_report/商汤科技深度研报_20250711_103844.md",
        "/Users/mac/Documents/ai/bishai/financial_research_report/4Paradigm深度研报_优化版_20250723_114029.md"
    ]
    
    for test_file in test_files:
        if os.path.exists(test_file):
            print(f"\n开始转换: {os.path.basename(test_file)}")
            result = convert_md_to_docx_pure_python(test_file)
            if result:
                print(f"✅ 转换成功！输出文件：{result}")
            else:
                print(f"❌ 转换失败：{test_file}")
        else:
            print(f"⚠️ 文件不存在: {test_file}")